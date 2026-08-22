# -*- coding: utf-8 -*-
"""生成第一问求解过程 Word 文档「解题过程.docx」（位于工作区根目录）。"""
import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
WORKSPACE = os.path.dirname(BASE)
OUT = os.path.join(WORKSPACE, "解题过程.docx")
IMG = os.path.join(BASE, "盘入示意图_计算结果.png")

doc = Document()

sec = doc.sections[0]
sec.page_width = Cm(21.0)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.0)
sec.bottom_margin = Cm(2.0)
sec.left_margin = Cm(2.0)
sec.right_margin = Cm(2.0)

def set_style_font(style, ascii_font, ea_font, size, bold=False):
    style.font.name = ascii_font
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), ascii_font)
    rfonts.set(qn('w:hAnsi'), ascii_font)
    rfonts.set(qn('w:eastAsia'), ea_font)

normal = doc.styles['Normal']
set_style_font(normal, 'Times New Roman', '宋体', 10.5)
normal.paragraph_format.space_after = Pt(3)
normal.paragraph_format.line_spacing = 1.25

set_style_font(doc.styles['Heading 1'], 'Times New Roman', '黑体', 14, bold=True)
doc.styles['Heading 1'].paragraph_format.space_before = Pt(10)
doc.styles['Heading 1'].paragraph_format.space_after = Pt(6)
set_style_font(doc.styles['Heading 2'], 'Times New Roman', '黑体', 12, bold=True)
doc.styles['Heading 2'].paragraph_format.space_before = Pt(6)
doc.styles['Heading 2'].paragraph_format.space_after = Pt(4)

TOKEN_RE = re.compile(r'(\^\{[^}]*\}|_\{[^}]*\})')

def add_marked_runs(par, text, ascii_font='Times New Roman', ea_font='宋体',
                    size=10.5, bold=False, italic=False):
    for token in TOKEN_RE.split(text):
        if not token:
            continue
        if token.startswith('^{'):
            r = par.add_run(token[2:-1])
            r.font.superscript = True
        elif token.startswith('_{'):
            r = par.add_run(token[2:-1])
            r.font.subscript = True
        else:
            r = par.add_run(token)
        r.font.name = ascii_font
        r._element.rPr.rFonts.set(qn('w:ascii'), ascii_font)
        r._element.rPr.rFonts.set(qn('w:hAnsi'), ascii_font)
        r._element.rPr.rFonts.set(qn('w:eastAsia'), ea_font)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
    return par

def body(text, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_marked_runs(p, text, size=10.5)
    return p

def bullet(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(21)
    add_marked_runs(p, text, size=10.5)
    return p

def eq(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    add_marked_runs(p, text, ascii_font='Cambria Math', ea_font='宋体', size=10.5)
    return p

def heading(text, level=1):
    p = doc.add_heading('', level=level)
    add_marked_runs(p, text, ascii_font='Times New Roman', ea_font='黑体',
                    size=14 if level == 1 else 12, bold=True)
    return p

def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    add_marked_runs(p, text, size=9, bold=True)
    return p

def shade(cell, fill='D9D9D9'):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def cell_text(cell, text, size=9, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    add_marked_runs(p, text, size=size, bold=bold)

def make_table(header, rows, widths_cm):
    t = doc.add_table(rows=len(rows) + 1, cols=len(header))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for j, h in enumerate(header):
        cell_text(t.rows[0].cells[j], h, bold=True)
        shade(t.rows[0].cells[j])
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell_text(t.rows[i].cells[j], val)
    for row in t.rows:
        for j, w in enumerate(widths_cm):
            row.cells[j].width = Cm(w)
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))
    return t

def set_bottom_border(p, color='999999'):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

# ---------- 标题 ----------
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
add_marked_runs(p, '2024年高教社杯全国大学生数学建模竞赛', ascii_font='Times New Roman', ea_font='黑体', size=14, bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
add_marked_runs(p, 'A题「板凳龙」第一问求解过程', ascii_font='Times New Roman', ea_font='黑体', size=16, bold=True)

# ---------- 1 问题重述 ----------
heading('1 问题重述')
body('某板凳龙由223节板凳组成：龙头板长341 cm，中间221节龙身板长220 cm，龙尾板长220 cm，板宽30 cm。每节板凳两端各有一个孔，孔心距板端27.5 cm。相邻板凳通过把手连接，因此同一块板上两个孔心的距离（即相邻把手中心距）为：')
bullet('· 龙头板：3.41 − 2×0.275 = 2.86 m；')
bullet('· 龙身与龙尾板：2.20 − 2×0.275 = 1.65 m。')
body('223节板凳共有224个把手中心：龙头前把手1个、相邻两板共用的连接把手222个、龙尾后把手1个。')
body('舞龙队沿螺距为55 cm的等距螺线顺时针盘入，各把手中心均位于螺线上；龙头前把手沿螺线行进的速度恒为1 m/s；初始时刻（t = 0）龙头位于螺线第16圈上的A点（题图4中，A点在x轴正半轴上）。要求计算0～300 s内每隔1 s全部224个把手的位置与速度，结果保留6位小数。')

# ---------- 2 模型假设 ----------
heading('2 模型假设')
body('(1) 把手视为质点：忽略孔径与板宽的影响（孔径与板宽只与问题2的碰撞判断有关）。')
body('(2) 板凳为刚性杆：同一块板两端的把手中心距在运动过程中恒等于相应的孔心距2.86 m或1.65 m。')
body('(3) 所有把手中心严格位于等距螺线上（题目给定）。')
body('(4) 龙头前把手沿螺线以恒定速率1 m/s行进，忽略起步的瞬态过程。')

# ---------- 3 符号说明 ----------
heading('3 符号说明')
caption('表1 主要符号说明')
sym_rows = [
    ('a', '螺线系数，a = p/(2π) ≈ 0.0875354 m/rad'),
    ('θ、r', '把手中心的极角（rad）与极径（m）'),
    ('P(θ)', '极角θ处螺线上点的直角坐标'),
    ('s(θ)', '螺线自极点（θ = 0）到θ的弧长（m）'),
    ('L_{i}', '第i段相邻把手中心距（m）：L_{1} = 2.86，其余L_{i} = 1.65'),
    ('v_{i}', '第i个把手沿螺线的速率（m/s）'),
]
make_table(['符号', '含义'], sym_rows, [3.2, 13.8])

# ---------- 4 坐标系与等距螺线方程 ----------
heading('4 坐标系与等距螺线方程')
body('以螺线中心（盘心）O为原点建立平面直角坐标系，取A点所在方向为x轴正方向。螺距p = 0.55 m的等距螺线（阿基米德螺线）的极坐标方程为')
eq('r(θ) = aθ，a = p/(2π) = 0.55/(2π) ≈ 0.0875354 m/rad')
body('直角坐标参数形式为')
eq('P(θ) = (x(θ), y(θ)) = (aθcosθ, aθsinθ)')
body('初始条件：龙头初始位于螺线第16圈、且A点在+x轴上，故龙头前把手初始极角θ_{0}(0) = 32π，初始极径r_{0} = a·32π = 0.55×16 = 8.8 m，即A = (8.800000, 0.000000)。顺时针盘入对应极角θ随时间减小。')
body('方向自洽性检验：龙身位于龙头走过的螺线外侧（θ增大的一侧），t = 0时龙尾后把手极角约136.2 rad（约第21.7圈）、极径约11.92 m，整条龙在螺线上占据的弧长约369.56 m，与223节板凳的板长总和369.16 m吻合，说明“龙身沿龙头来路向螺线外侧延伸”的方向选择自洽。')

# ---------- 5 螺线弧长公式 ----------
heading('5 螺线弧长公式')
body('由极坐标弧微分 ds² = dr² + (r dθ)² 及 r′ = dr/dθ = a，得')
eq('ds/dθ = √(r² + r′²) = a√(θ² + 1)')
body('自极点（θ = 0）到θ的弧长为')
eq('s(θ) = a∫_{0}^{θ} √(u² + 1) du = (a/2)[θ√(θ²+1) + arsinh θ]')
body('其中 arsinh θ = ln(θ + √(θ²+1)) 为反双曲正弦。例如 s(32π) ≈ 442.590256 m。')

# ---------- 6 龙头前把手运动方程 ----------
heading('6 龙头前把手运动方程')
body('龙头前把手沿螺线以恒定速率1 m/s顺时针行进（θ减小），由弧长对时间的变化率可得')
eq('v = ds/dt = a√(θ_{0}²+1)·|dθ_{0}/dt| = 1 ⇒ dθ_{0}/dt = −1/√(r_{0}²+a²) < 0')
body('等价地，沿螺线弧长线性推进：')
eq('s(θ_{0}(t)) = s(32π) − t，t ∈ [0, 300]')
body('对每个整数时刻t，用Newton迭代反解θ_{0}(t)：')
eq('θ ← θ − [s(θ) − (s(32π) − t)] / [a√(θ²+1)]')
body('初值取θ ≈ √((32π)² − 2t/a)，迭代至弧长残差小于1×10^{−15} m（约4～5次迭代）。t = 300 s时θ_{0} ≈ 57.032 rad（约第9.1圈），r ≈ 4.99 m，龙头仍在螺线上。')

# ---------- 7 相邻把手递推 ----------
heading('7 相邻把手递推：弦长精确求解（刚性杆约束）')
body('记把手i（i = 0为龙头前把手，i = 1, …, 221为第i节龙身前把手，i = 222为龙尾前把手，i = 223为龙尾后把手），其极角为θ_{i}。相邻把手中心距为')
eq('L_{1} = 2.86 m（龙头板）；L_{i} = 1.65 m（i = 2, …, 223）')
body('龙身位于龙头走过的螺线外侧（θ更大的一侧），对i = 1, …, 223依次求解一维方程')
eq('F(θ_{i}) = |P(θ_{i}) − P(θ_{i−1})|² − L_{i}² = 0，θ_{i} > θ_{i−1}')
body('采用Newton迭代（必要时用二分法兜底）：')
eq('θ_{i} ← θ_{i} − F/F′，F′ = 2[P(θ_{i}) − P(θ_{i−1})]·P′(θ_{i})')
body('其中P′(θ) = a(cosθ − θsinθ, sinθ + θcosθ)。初值取弧长近似θ_{i} ≈ θ_{i−1} + L_{i}/√(r_{i−1}²+a²)，每节约4次迭代收敛。求解后224个把手全部严格落在螺线上，且相邻把手弦长严格等于板长（最大残差约1.7×10^{−13} m）。')
body('弦长模型与弧长间距模型对比：若近似为相邻把手沿螺线弧长相距L_{i}，则弦长略小于弧长，单节误差约为L_{i}³/(24R²)（R为曲率半径）。当R ≈ 8.8 m时每节约差2.4 mm，222节累计使龙尾位置偏差约0.40 m（t = 0）至0.86 m（t = 300）。为严格满足“板凳为刚性杆”的假设，本解采用弦长精确模型。')

# ---------- 8 速度递推 ----------
heading('8 速度递推：弦长约束的隐式求导')
body('对约束|P(θ_{i}) − P(θ_{i−1})| = L_{i}关于时间t求导。设V_{i} = dP(θ_{i})/dt为把手i的速度向量，刚性杆两端沿杆方向的速率分量相等：')
eq('[P(θ_{i}) − P(θ_{i−1})]·(V_{i} − V_{i−1}) = 0')
body('把手沿螺线运动，速度方向沿切线：V_{i} = v_{i}e_{i}，其中顺时针单位切向量为')
eq('e_{i} = −P′(θ_{i})/|P′(θ_{i})| = −P′(θ_{i})/√(r_{i}²+a²)')
body('记单位杆向link = (P(θ_{i}) − P(θ_{i−1}))/L_{i}，代入得速度递推公式')
eq('v_{i} = v_{i−1}·(link·e_{i−1})/(link·e_{i})，v_{0} = 1 m/s')
body('等价的极角形式：θ_{i}′ = θ_{i−1}′·[(P(θ_{i−1}) − P(θ_{i}))·P′(θ_{i−1})] / [(P(θ_{i−1}) − P(θ_{i}))·P′(θ_{i})]，v_{i} = √(r_{i}²+a²)·|θ_{i}′|。')
body('由于龙身的跟随滞后，越靠龙尾的把手速率略小于1 m/s：t = 0时龙头后把手（第1节龙身前把手）速率为0.999971 m/s，龙尾后把手为0.999311 m/s；随盘入半径变小，龙尾后把手速率降至t = 300 s的0.996478 m/s。')

# ---------- 9 数值算法与精度控制 ----------
heading('9 数值算法与精度控制')
body('对t = 0, 1, …, 300共301个时刻依次执行：')
body('步骤1：按第6节弧长反解龙头极角θ_{0}(t)（Newton迭代，弧长误差小于1×10^{−15} m）；')
body('步骤2：按第7节逐节弦长递推224个θ_{i}（Newton迭代，弦长残差小于1×10^{−12} m）；')
body('步骤3：按第8节隐式求导递推224个速率v_{i}。')
body('全部采用双精度浮点运算，无随机与拟合成分；总计算量约301×224次一维求根，秒级完成。')

# ---------- 10 数值校验 ----------
heading('10 数值校验')
caption('表2 数值校验结果')
val_rows = [
    ('相邻把手弦长最大残差', '1.69×10^{−13} m'),
    ('龙头0→300 s弧长减少量', '300.000000000 m（理论值300 m）'),
    ('速度递推与中心有限差分最大相对误差', '6.8×10^{−8}（t = 100 s，步长1×10^{−5} s）'),
    ('t = 0龙头位置', '(8.800000000, 0.000000000)，与题图4的A点一致'),
    ('t = 60 s龙头位置', '(5.799209030, −5.771092341)，y < 0与顺时针盘入方向一致'),
    ('弦长模型与弧长间距模型最大位置偏差', 't = 0为0.403 m，t = 300为0.860 m'),
    ('与工作区独立复现参考值最大绝对偏差', '3.1×10^{−13}'),
]
make_table(['校验项目', '校验结果'], val_rows, [6.0, 11.0])

# ---------- 11 结果 ----------
heading('11 结果')
body('按照题目要求，表3给出0、60、120、180、240、300 s时刻龙头、第1、51、101、151、201节龙身前把手及龙尾后把手的位置，表4给出对应速度。')

heading('11.1 表3：把手位置（单位m，保留6位小数）', 2)
caption('表3 把手位置（单位m，保留6位小数）')
pos_rows = [
    ('龙头 x', '8.800000', '5.799209', '-4.084887', '-2.963609', '2.594494', '4.420274'),
    ('龙头 y', '0.000000', '-5.771092', '-6.304479', '6.094780', '-5.356743', '2.320429'),
    ('第1节龙身 x', '8.363824', '7.456758', '-1.445473', '-5.237118', '4.821221', '2.459489'),
    ('第1节龙身 y', '2.826544', '-3.440399', '-7.405883', '4.359627', '-3.561949', '4.402476'),
    ('第51节龙身 x', '-9.518732', '-8.686317', '-5.543150', '2.890455', '5.980011', '-6.301346'),
    ('第51节龙身 y', '1.341137', '2.540108', '6.377946', '7.249289', '-3.827758', '0.465829'),
    ('第101节龙身 x', '2.913983', '5.687116', '5.361939', '1.898794', '-4.917371', '-6.237722'),
    ('第101节龙身 y', '-9.918311', '-8.001384', '-7.557638', '-8.471614', '-6.379874', '3.936008'),
    ('第151节龙身 x', '10.861726', '6.682311', '2.388757', '1.005154', '2.965378', '7.040740'),
    ('第151节龙身 y', '1.828753', '8.134544', '9.727411', '9.424751', '8.399721', '4.393013'),
    ('第201节龙身 x', '4.555102', '-6.619664', '-10.627211', '-9.287720', '-7.457151', '-7.458662'),
    ('第201节龙身 y', '10.725118', '9.025570', '1.359847', '-4.246673', '-6.180726', '-5.263384'),
    ('龙尾（后） x', '-5.305444', '7.364557', '10.974348', '7.383896', '3.241051', '1.785033'),
    ('龙尾（后） y', '-10.676584', '-8.797992', '0.843473', '7.492370', '9.469336', '9.301164'),
]
make_table(['部位', '0 s', '60 s', '120 s', '180 s', '240 s', '300 s'], pos_rows, [2.3, 2.45, 2.45, 2.45, 2.45, 2.45, 2.45])

heading('11.2 表4：把手速度（单位m/s，保留6位小数）', 2)
caption('表4 把手速度（单位m/s，保留6位小数）')
spd_rows = [
    ('龙头', '1.000000', '1.000000', '1.000000', '1.000000', '1.000000', '1.000000'),
    ('第1节龙身', '0.999971', '0.999961', '0.999945', '0.999917', '0.999859', '0.999709'),
    ('第51节龙身', '0.999742', '0.999662', '0.999538', '0.999331', '0.998941', '0.998065'),
    ('第101节龙身', '0.999575', '0.999453', '0.999269', '0.998971', '0.998435', '0.997302'),
    ('第151节龙身', '0.999448', '0.999299', '0.999078', '0.998727', '0.998115', '0.996861'),
    ('第201节龙身', '0.999348', '0.999180', '0.998935', '0.998551', '0.997894', '0.996574'),
    ('龙尾（后）', '0.999311', '0.999136', '0.998883', '0.998489', '0.997816', '0.996478'),
]
make_table(['部位', '0 s', '60 s', '120 s', '180 s', '240 s', '300 s'], spd_rows, [2.3, 2.45, 2.45, 2.45, 2.45, 2.45, 2.45])
body('全量数据（301个时刻×224个把手的位置与速度）保存在配套文件result1.xlsx的“位置”与“速度”工作表中。')

# ---------- 12 龙形示意图 ----------
heading('12 龙形示意图')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(IMG, width=Cm(14))
caption('图1  0 s与300 s板凳龙盘入形态对比示意图')

# ---------- 13 程序实现 ----------
heading('13 程序实现')
body('求解程序solve_q1.py仅依赖标准库math与openpyxl，在Windows PowerShell中运行：')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Pt(21)
code_lines = [
    "$env:PYTHONIOENCODING='utf-8'",
    'cd 第一问求解',
    'C:\\ProgramData\\Anaconda3\\python.exe solve_q1.py',
]
for i, code_line in enumerate(code_lines):
    if i > 0:
        p.add_run().add_break()
    r = p.add_run(code_line)
    r.font.name = 'Consolas'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r.font.size = Pt(9)
body('运行后生成result1.xlsx与表3、表4的文本结果。')

# ---------- 页眉与页脚 ----------
hp = sec.header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_marked_runs(hp, 'A题「板凳龙」第一问求解过程', ascii_font='Times New Roman', ea_font='宋体', size=9)
set_bottom_border(hp)

fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run('第 ')
r.font.name = 'Times New Roman'
r.font.size = Pt(9)
r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run = fp.add_run()
fld1 = OxmlElement('w:fldChar')
fld1.set(qn('w:fldCharType'), 'begin')
instr = OxmlElement('w:instrText')
instr.set(qn('xml:space'), 'preserve')
instr.text = 'PAGE'
fld2 = OxmlElement('w:fldChar')
fld2.set(qn('w:fldCharType'), 'end')
run._r.append(fld1)
run._r.append(instr)
run._r.append(fld2)
run.font.name = 'Times New Roman'
run.font.size = Pt(9)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
r = fp.add_run(' 页')
r.font.name = 'Times New Roman'
r.font.size = Pt(9)
r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ---------- 属性与保存 ----------
cp = doc.core_properties
cp.title = '2024年高教社杯全国大学生数学建模竞赛 A题「板凳龙」第一问求解过程'
cp.author = ''
cp.comments = ''

doc.save(OUT)
print('saved:', OUT)
